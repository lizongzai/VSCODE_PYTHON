from docx import Document
from docx.document import Document as Doc

# 将真实信息用字典的方式保存在列表中
employees = [
   {
        'name': '骆昊',
        'id': '100200198011280001',
        'sdate': '2008年3月1日',
        'edate': '2012年2月29日',
        'department': '产品研发',
        'position': '架构师',
        'company': '成都华为技术有限公司'
    },
    {
        'name': '王大锤',
        'id': '510210199012125566',
        'sdate': '2019年1月1日',
        'edate': '2021年4月30日',
        'department': '产品研发',
        'position': 'Python开发工程师',
        'company': '成都谷道科技有限公司'
    },
    {
        'name': '李元芳',
        'id': '2102101995103221599',
        'sdate': '2020年5月10日',
        'edate': '2021年3月5日',
        'department': '产品研发',
        'position': 'Java开发工程师',
        'company': '同城企业管理集团有限公司'
    }
]

# # 对列表进行循环遍历，批量生成Word文档 
# for emp_dict in employees:
#     # 读取离职证明模板文件
#     document = Document('resources/离职证明模板.docx')
#     # 循环遍历所有段落寻找占位符
#     for paragraphs in document.paragraphs:
#         if '{' not in paragraphs.text:
#             continue
#             # 不能直接修改段落内容，否则会丢失样式
#             # 所以需要对段落中的元素进行遍历并进行查找替换
#         for run in paragraphs.runs:
#             if '{' not in paragraphs.text:
#                 continue
#             # 将占位符换成实际内容
#             start, end = run.text.find('{'), run.text.find('}')
#             key, place_holder = run.text[start + 1:end], run.text[start:end + 1]
#             print(f"key='{key}' place_holder='{place_holder}'")

#             # run.text = run.text.replace(place_holder, emp_dict[key])
            
#     # 每个人对应保存一个Word文档
#     document.save(f'{emp_dict["name"]}离职证明.docx')            

for emp_dict in employees:
    document = Document('resources/离职证明模板.docx')
    for paragraph in document.paragraphs:
        text = paragraph.text
        # 遍历字典里的 key，替换占位符
        for key, value in emp_dict.items():
            text = text.replace(f'{{{key}}}', value)
        paragraph.text = text
    document.save(f'{emp_dict["name"]}离职证明.docx')
