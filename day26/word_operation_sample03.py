from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.document import Document as Doc

# -----------------------------
# 设置字体的辅助函数
# -----------------------------
def set_font(run, font_name='宋体', font_size=None, bold=False, underline=False):
    """确保 run 设置中文字体（EastAsia），避免口口问题"""
    run.font.name = font_name  # 西文字体
    r = run._element

    # 确保 rPr 存在
    if r.rPr is None:
        r.get_or_add_rPr()

    # 确保 rFonts 存在
    rFonts = r.rPr.rFonts
    if rFonts is None:
        rFonts = r.rPr.get_or_add_rFonts()

    rFonts.set(qn('w:eastAsia'), font_name)

    # 设置字体大小
    if font_size:
        run.font.size = font_size
    # 设置加粗
    run.bold = bold
    # 设置下划线
    run.underline = underline

# -----------------------------
# 创建文档
# -----------------------------
document = Document()  # type: Doc

# 添加大标题
title = document.add_heading('快快乐乐学Python', 0)
for run in title.runs:
    set_font(run, '微软雅黑', Pt(24), bold=True)

# 添加作者
author_paragraph = document.add_paragraph()
author_run = author_paragraph.add_run('作者：李小婕')
set_font(author_run, '微软雅黑', Pt(12))

# 添加段落
paragraph = document.add_paragraph('Python是一门非常流行的编程语言，它')
for run in paragraph.runs:
    set_font(run, '宋体', Pt(12))

# 添加加粗的“简单”
run = paragraph.add_run('简单')
set_font(run, '宋体', Pt(18), bold=True)

# 继续添加文本
run = paragraph.add_run('而且')
set_font(run, '宋体', Pt(12))

# 添加下划线的“优雅”
run = paragraph.add_run('优雅')
set_font(run, '宋体', Pt(18), underline=True)

# 继续添加句号
run = paragraph.add_run('。')
set_font(run, '宋体', Pt(12))

# 添加一级标题
heading1 = document.add_heading('Heading, level 1', level=1)
for run in heading1.runs:
    set_font(run, '宋体', Pt(16), bold=True)

# 添加带样式的段落
quote = document.add_paragraph('Intense quote', style='Intense Quote')
for run in quote.runs:
    set_font(run, '宋体', Pt(12))

# 添加无序列表
ul1 = document.add_paragraph('first item in unordered list', style='List Bullet')
for run in ul1.runs:
    set_font(run, '宋体', Pt(12))
ul2 = document.add_paragraph('second item in unordered list', style='List Bullet')
for run in ul2.runs:
    set_font(run, '宋体', Pt(12))

# 添加有序列表
ol1 = document.add_paragraph('first item in ordered list', style='List Number')
for run in ol1.runs:
    set_font(run, '宋体', Pt(12))
ol2 = document.add_paragraph('second item in ordered list', style='List Number')
for run in ol2.runs:
    set_font(run, '宋体', Pt(12))

# 添加图片（确保路径存在）
document.add_picture('resources/tut.jpg', width=Cm(5.2))

# 添加分页符
document.add_page_break()

# -----------------------------
# 添加表格
# -----------------------------
records = (
    ('骆昊', '男', '1995-5-5'),
    ('孙美丽', '女', '1992-2-2')
)

table = document.add_table(rows=1, cols=3)
table.style = 'Dark List'

# 添加表头
hdr_cells = table.rows[0].cells
headers = ['姓名', '性别', '出生日期']
for i, header in enumerate(headers):
    hdr_cells[i].text = header
    for p in hdr_cells[i].paragraphs:
        for run in p.runs:
            set_font(run, '宋体', Pt(12), bold=True)

# 添加数据行
for name, sex, birthday in records:
    row_cells = table.add_row().cells
    values = [name, sex, birthday]
    for i, value in enumerate(values):
        row_cells[i].text = value
        for p in row_cells[i].paragraphs:
            for run in p.runs:
                set_font(run, '宋体', Pt(12))

# 添加分页符
document.add_page_break()

# -----------------------------
# 保存文档
# -----------------------------
document.save('demo1.docx')
print("文档生成成功！")
