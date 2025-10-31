"""
Version: 1.0
Author: ChatGPT
Description: 我们可以使用如下所示的代码来生成一个简单的 Word 文档，并在其中进行一些基本的文字操作。
Date: 2025-10-31

Keyword arguments:
argument -- description
Return: return_description
"""
from docx import Document
from docx.shared import Pt, Cm

# 创建一个新的 Word 文档
document = Document()
# 添加标题
document.add_heading('Word 文档操作示例', level=1)
# 添加段落
paragraph = document.add_paragraph('这是一个使用 python-docx 库创建的 Word 文档示例。')
# 设置段落字体和大小
run = paragraph.runs[0]
run.font.size = Pt(12)
run.font.name = 'Arial'
run.font.bold = True
# 添加另一个段落
document.add_paragraph('我们可以在文档中添加多个段落，并设置不同的样式。')
# 添加一个带有自定义样式的段落
styled_paragraph = document.add_paragraph()
styled_run = styled_paragraph.add_run('这是一个自定义样式的段落，字体大小为14磅，行距为1.5倍。')
styled_run.font.size = Pt(14)
styled_paragraph.paragraph_format.line_spacing = 1.5
# 添加一个表格
table = document.add_table(rows=3, cols=3)
table.style = 'Table Grid'
# 填充表格数据
for row in table.rows:
    for cell in row.cells:
        cell.text = '单元格'
        cell.width = Cm(3)
# 保存文档
document.save('word_operation_sample01.docx')
