from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.oxml.ns import qn

from docx.document import Document as Doc

# 设置字体的辅助函数
def set_font(run, font_name='宋体'):
    """确保 run 设置中文字体（EastAsia），避免口口问题"""
    run.font.name = font_name  # 西文字体
    r = run._element

    # 确保 rPr 存在
    if r.rPr is None:
        r.get_or_add_rPr()

    # 确保 rFonts 存在
    rFonts = r.rPr.rFonts
    if rFonts is None:
        rFonts = r.rPr.get_or_add_rFonts()

    rFonts.set(qn('w:eastAsia'), font_name)


# 创建代表Word文档的Doc对象
document = Document()  # type: Doc
# 添加大标题
document.add_heading('快快乐乐学Python', 0)

# 添加作者
set_font_run = document.add_paragraph().add_run('作者：李小婕')
# 设定字体,微软雅黑。
set_font(set_font_run, '微软雅黑')

# 添加段落
paragraph = document.add_paragraph('Python是一门非常流行的编程语言，它')
set_font(paragraph.runs[0], '宋体')

# 添加加粗的“简单”二字
run = paragraph.add_run('简单')
set_font(run, '宋体') # 设置字体为宋体,否则输出会出现问题

# 设置字体为粗体
run.bold = True
# 设置字体大小为18磅
run.font.size = Pt(18)
# 继续添加文本
paragraph.add_run('而且')
# 添加下划线的“优雅”二字
run = paragraph.add_run('优雅')
set_font(run, '宋体')

# 设置字体为下划线
run.font.size = Pt(18)
# 设置字体为下划线
run.underline = True
# 继续添加文本
paragraph.add_run('。')

# 添加一级标题
document.add_heading('Heading, level 1', level=1)
# 添加带样式的段落
document.add_paragraph('Intense quote', style='Intense Quote')

# 添加无序列表
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'second item in ordered list', style='List Bullet'
)

# 添加有序列表
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
document.add_paragraph(
    'second item in ordered list', style='List Number'
)

# 添加图片（注意路径和图片必须要存在）
document.add_picture('resources/tut.jpg', width=Cm(5.2))

# 添加分节符
# document.add_section()

# # 添加分节符（改成分页符）
# document.add_page_break()


# 获取图片所在段落
last_paragraph = document.paragraphs[-1]
# 设置段落不分页，和下一段（表格）保持在同一页
last_paragraph.paragraph_format.keep_with_next = True


records = (
    ('骆昊', '男', '1995-5-5'),
    ('孙美丽', '女', '1992-2-2')
)

# 添加表格
table = document.add_table(rows=1, cols=3)
table.style = 'Dark List'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '姓名'
set_font(hdr_cells[0].paragraphs[0].runs[0], '宋体')
hdr_cells[1].text = '性别'
set_font(hdr_cells[1].paragraphs[0].runs[0], '宋体')
hdr_cells[2].text = '出生日期'
set_font(hdr_cells[2].paragraphs[0].runs[0], '宋体')


# 为表格添加行
for name, sex, birthday in records:
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = sex
    row_cells[2].text = birthday
    set_font(row_cells[0].paragraphs[0].runs[0], '宋体')
    set_font(row_cells[1].paragraphs[0].runs[0], '宋体')
    set_font(row_cells[2].paragraphs[0].runs[0], '宋体')

# # 添加分页符
document.add_page_break()

# 保存文档
document.save('demo.docx')